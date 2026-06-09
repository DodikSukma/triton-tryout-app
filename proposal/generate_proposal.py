from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2.5)

# ── Default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ── Helper: set paragraph font
def para_font(para, size=None, bold=False, color=None, name='Calibri'):
    for run in para.runs:
        run.font.name  = name
        if size:  run.font.size  = Pt(size)
        if bold:  run.font.bold  = True
        if color: run.font.color.rgb = RGBColor(*color)

# ── Helper: set cell bg color
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

# ── Helper: add thick border to table
def set_table_style(table):
    tbl   = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top','left','bottom','right','insideH','insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'),   'single')
        border.set(qn('w:sz'),    '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'D1D5DB')
        tblBorders.append(border)
    tblPr.append(tblBorders)

# ══════════════════════════════════════════════════════════
# HEADER COMPANY
# ══════════════════════════════════════════════════════════
header_para = doc.add_paragraph()
header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = header_para.add_run('ATALA PROJECT')
run.font.name  = 'Calibri'
run.font.size  = Pt(20)
run.font.bold  = True
run.font.color.rgb = RGBColor(30, 64, 175)   # blue-800

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = sub.add_run('Technology Solutions & Digital Infrastructure')
r.font.name  = 'Calibri'
r.font.size  = Pt(10)
r.font.color.rgb = RGBColor(100, 116, 139)   # slate-500

doc.add_paragraph()

# ── Divider line
div = doc.add_paragraph()
div_run = div.add_run('─' * 85)
div_run.font.color.rgb = RGBColor(30, 64, 175)
div_run.font.size = Pt(9)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════
# PROPOSAL TITLE
# ══════════════════════════════════════════════════════════
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('PROPOSAL LAYANAN INFRASTRUKTUR CLOUD')
r.font.name  = 'Calibri'
r.font.size  = Pt(18)
r.font.bold  = True
r.font.color.rgb = RGBColor(15, 23, 42)   # slate-900

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = subtitle.add_run('Platform Tryout Online — Triton Denpasar')
r.font.name  = 'Calibri'
r.font.size  = Pt(12)
r.font.color.rgb = RGBColor(71, 85, 105)

doc.add_paragraph()

# ── Info box (Kepada / Dari / Tanggal / No)
today = datetime.date.today()
info_lines = [
    ('Kepada',    'TRITON BIMBEL'),
    ('Dari',      'Atala Project'),
    ('Tanggal',   today.strftime('%d %B %Y')),
    ('No. Proposal', f'ATP-{today.strftime("%Y%m")}-001'),
    ('Validitas', '30 hari dari tanggal proposal'),
]
info_table = doc.add_table(rows=len(info_lines), cols=3)
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, (label, value) in enumerate(info_lines):
    row = info_table.rows[i]
    row.cells[0].text = label
    row.cells[1].text = ':'
    row.cells[2].text = value
    for j, cell in enumerate(row.cells):
        p = cell.paragraphs[0]
        for run in p.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(10.5)
            if j == 0: run.font.bold = True
            if j == 2: run.font.color.rgb = RGBColor(30, 64, 175)

doc.add_paragraph()
div2 = doc.add_paragraph()
div_run2 = div2.add_run('─' * 85)
div_run2.font.color.rgb = RGBColor(30, 64, 175)
div_run2.font.size = Pt(9)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════
# 1. KATA PENGANTAR
# ══════════════════════════════════════════════════════════
def add_section_title(doc, number, title_text):
    p = doc.add_paragraph()
    r = p.add_run(f'{number}. {title_text}')
    r.font.name  = 'Calibri'
    r.font.size  = Pt(13)
    r.font.bold  = True
    r.font.color.rgb = RGBColor(30, 64, 175)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)

add_section_title(doc, '1', 'KATA PENGANTAR')

pengantar = doc.add_paragraph(
    'Dengan hormat,\n\n'
    'Atala Project mengucapkan terima kasih atas kepercayaan TRITON BIMBEL dalam mempertimbangkan '
    'layanan infrastruktur digital kami. Proposal ini disusun sebagai respons atas kebutuhan platform '
    'tryout online dengan kapasitas ±500 siswa yang dapat mengakses sistem secara bersamaan.\n\n'
    'Kami memahami bahwa keandalan sistem adalah prioritas utama, khususnya saat sesi tryout berlangsung. '
    'Oleh karena itu, kami menyiapkan dua pilihan paket yang dapat disesuaikan dengan anggaran dan '
    'kebutuhan operasional TRITON BIMBEL.'
)
pengantar.paragraph_format.space_after = Pt(4)
for run in pengantar.runs:
    run.font.name = 'Calibri'
    run.font.size = Pt(11)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════
# 2. LINGKUP LAYANAN
# ══════════════════════════════════════════════════════════
add_section_title(doc, '2', 'LINGKUP LAYANAN')

scope_items = [
    '✅  Penyediaan dan pengelolaan server cloud (VPS) dedicated',
    '✅  Deployment platform tryout online Triton Denpasar',
    '✅  Konfigurasi database PostgreSQL (4 database terpisah)',
    '✅  Setup API Gateway + Microservices (Auth, User, Soal, Jawaban)',
    '✅  Konfigurasi keamanan: HTTPS, cookie HttpOnly, session management',
    '✅  Monitoring uptime dan performa server 24/7',
    '✅  Backup data otomatis mingguan',
    '✅  Technical support via WhatsApp (response time maks. 24 jam)',
    '✅  Training penggunaan panel admin untuk staff TRITON BIMBEL',
]

for item in scope_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(2)
    for run in p.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(10.5)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════
# 3. ANALISIS KEBUTUHAN
# ══════════════════════════════════════════════════════════
add_section_title(doc, '3', 'ANALISIS KEBUTUHAN TEKNIS')

analisis = doc.add_paragraph(
    'Berdasarkan skenario operasional platform tryout dengan ±500 siswa secara bersamaan, '
    'berikut adalah profil beban sistem yang kami identifikasi:'
)
for run in analisis.runs:
    run.font.name = 'Calibri'
    run.font.size = Pt(11)

doc.add_paragraph()

# Analisis table
analisis_headers = ['Skenario', 'Estimasi Beban', 'Keterangan']
analisis_data = [
    ['Siswa login serentak',      '~500 req dalam 5 menit',   'Beban ringan, tersebar'],
    ['Selama ujian berlangsung',  '~4 req/detik (auto-save)', 'Steady, terkelola'],
    ['Semua submit bersamaan',    '~500 req dalam 30 detik',  'Peak load — titik kritis'],
    ['Hitung nilai & simpan DB',  '~500 query berat serentak','CPU & DB intensive'],
]

tbl_a = doc.add_table(rows=1 + len(analisis_data), cols=3)
tbl_a.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_style(tbl_a)

# Header row
hdr_cells = tbl_a.rows[0].cells
for i, h in enumerate(analisis_headers):
    hdr_cells[i].text = h
    set_cell_bg(hdr_cells[i], '1E40AF')
    for run in hdr_cells[i].paragraphs[0].runs:
        run.font.bold  = True
        run.font.color.rgb = RGBColor(255,255,255)
        run.font.name  = 'Calibri'
        run.font.size  = Pt(10.5)

# Data rows
for r_i, row_data in enumerate(analisis_data):
    row = tbl_a.rows[r_i + 1]
    for c_i, val in enumerate(row_data):
        row.cells[c_i].text = val
        bg = 'F8FAFC' if r_i % 2 == 0 else 'FFFFFF'
        set_cell_bg(row.cells[c_i], bg)
        for run in row.cells[c_i].paragraphs[0].runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(10)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════
# 4. PAKET LAYANAN
# ══════════════════════════════════════════════════════════
add_section_title(doc, '4', 'PAKET LAYANAN')

# ── 4a. Tier 1
p = doc.add_paragraph()
r = p.add_run('📦  PAKET STARTER — Tier 1')
r.font.name  = 'Calibri'
r.font.size  = Pt(12)
r.font.bold  = True
r.font.color.rgb = RGBColor(30, 64, 175)
p.paragraph_format.space_before = Pt(8)

tier1_spec = doc.add_paragraph(
    'Spesifikasi Server: 4 vCPU · 8 GB RAM · 80 GB SSD NVMe · Bandwidth 2 TB/bln'
)
for run in tier1_spec.runs:
    run.font.name  = 'Calibri'
    run.font.size  = Pt(10.5)
    run.font.color.rgb = RGBColor(71, 85, 105)

doc.add_paragraph()

# Tier 1 table
t1_headers = ['Komponen', 'Detail']
t1_data = [
    ['Setup & Deployment Fee',   'Rp 5.000.000  (dibayar sekali)'],
    ['Biaya Bulanan',            'Rp 1.500.000 / bulan'],
    ['Total Tahun Pertama',      'Rp 23.000.000'],
    ['Perpanjangan Tahun ke-2+', 'Rp 18.000.000 / tahun'],
    ['Kapasitas Siswa',          'Optimal s/d 300 siswa concurrent'],
    ['Uptime Guarantee',         '99% per bulan'],
]

tbl1 = doc.add_table(rows=1 + len(t1_data), cols=2)
tbl1.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_style(tbl1)

hdr1 = tbl1.rows[0].cells
for i, h in enumerate(t1_headers):
    hdr1[i].text = h
    set_cell_bg(hdr1[i], '3B82F6')
    for run in hdr1[i].paragraphs[0].runs:
        run.font.bold  = True
        run.font.color.rgb = RGBColor(255,255,255)
        run.font.name  = 'Calibri'
        run.font.size  = Pt(10.5)

for r_i, row_data in enumerate(t1_data):
    row = tbl1.rows[r_i + 1]
    for c_i, val in enumerate(row_data):
        row.cells[c_i].text = val
        bg = 'EFF6FF' if r_i % 2 == 0 else 'FFFFFF'
        set_cell_bg(row.cells[c_i], bg)
        for run in row.cells[c_i].paragraphs[0].runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(10.5)
            if c_i == 0: run.font.bold = True
            if val.startswith('Rp') and 'Total' in t1_data[r_i][0]:
                run.font.bold  = True
                run.font.color.rgb = RGBColor(30, 64, 175)

doc.add_paragraph()

# Keunggulan & Kelemahan Tier 1
ku1_headers = ['✅ Keunggulan', '⚠️ Keterbatasan']
ku1_data = [
    ['Biaya lebih terjangkau',            'Berpotensi lambat saat 500 siswa submit bersamaan'],
    ['Cocok untuk tryout terjadwal rapi', 'RAM lebih terbatas untuk proses scoring massal'],
    ['Cukup untuk skala awal operasional','Tidak ideal jika tryout sering berjalan paralel'],
]

tbl_ku1 = doc.add_table(rows=1 + len(ku1_data), cols=2)
tbl_ku1.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_style(tbl_ku1)

for i, h in enumerate(ku1_headers):
    tbl_ku1.rows[0].cells[i].text = h
    bg = '166534' if i == 0 else '991B1B'
    set_cell_bg(tbl_ku1.rows[0].cells[i], bg)
    for run in tbl_ku1.rows[0].cells[i].paragraphs[0].runs:
        run.font.bold  = True
        run.font.color.rgb = RGBColor(255,255,255)
        run.font.name  = 'Calibri'
        run.font.size  = Pt(10.5)

for r_i, (pro, con) in enumerate(ku1_data):
    row = tbl_ku1.rows[r_i + 1]
    row.cells[0].text = pro
    row.cells[1].text = con
    bg = 'F0FDF4' if r_i % 2 == 0 else 'FFFFFF'
    set_cell_bg(row.cells[0], bg)
    bg2 = 'FEF2F2' if r_i % 2 == 0 else 'FFFFFF'
    set_cell_bg(row.cells[1], bg2)
    for c_i in [0, 1]:
        for run in row.cells[c_i].paragraphs[0].runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(10)

doc.add_paragraph()

# ── 4b. Tier 2
p2 = doc.add_paragraph()
r2 = p2.add_run('🏆  PAKET PROFESIONAL — Tier 2  (Rekomendasi)')
r2.font.name  = 'Calibri'
r2.font.size  = Pt(12)
r2.font.bold  = True
r2.font.color.rgb = RGBColor(5, 150, 105)   # emerald
p2.paragraph_format.space_before = Pt(12)

tier2_spec = doc.add_paragraph(
    'Spesifikasi Server: 6 vCPU · 16 GB RAM · 160 GB SSD NVMe · Bandwidth 4 TB/bln'
)
for run in tier2_spec.runs:
    run.font.name  = 'Calibri'
    run.font.size  = Pt(10.5)
    run.font.color.rgb = RGBColor(71, 85, 105)

doc.add_paragraph()

t2_data = [
    ['Setup & Deployment Fee',   'Rp 7.500.000  (dibayar sekali)'],
    ['Biaya Bulanan',            'Rp 2.500.000 / bulan'],
    ['Total Tahun Pertama',      'Rp 37.500.000'],
    ['Perpanjangan Tahun ke-2+', 'Rp 30.000.000 / tahun'],
    ['Kapasitas Siswa',          'Optimal s/d 500+ siswa concurrent'],
    ['Uptime Guarantee',         '99.5% per bulan'],
]

tbl2 = doc.add_table(rows=1 + len(t2_data), cols=2)
tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_style(tbl2)

hdr2 = tbl2.rows[0].cells
for i, h in enumerate(t1_headers):
    hdr2[i].text = h
    set_cell_bg(hdr2[i], '059660')
    for run in hdr2[i].paragraphs[0].runs:
        run.font.bold  = True
        run.font.color.rgb = RGBColor(255,255,255)
        run.font.name  = 'Calibri'
        run.font.size  = Pt(10.5)

for r_i, row_data in enumerate(t2_data):
    row = tbl2.rows[r_i + 1]
    for c_i, val in enumerate(row_data):
        row.cells[c_i].text = val
        bg = 'ECFDF5' if r_i % 2 == 0 else 'FFFFFF'
        set_cell_bg(row.cells[c_i], bg)
        for run in row.cells[c_i].paragraphs[0].runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(10.5)
            if c_i == 0: run.font.bold = True
            if 'Total' in t2_data[r_i][0]:
                run.font.bold  = True
                run.font.color.rgb = RGBColor(5, 150, 105)

doc.add_paragraph()

ku2_data = [
    ['Nyaman untuk 500 siswa submit bersamaan', 'Biaya lebih tinggi dari Tier 1'],
    ['16 GB RAM — scoring massal tanpa hambatan','—'],
    ['Ruang tumbuh: bisa tambah 200-300 siswa lagi','—'],
    ['Performa stabil sepanjang sesi tryout','—'],
]

tbl_ku2 = doc.add_table(rows=1 + len(ku2_data), cols=2)
tbl_ku2.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_style(tbl_ku2)

for i, h in enumerate(ku1_headers):
    tbl_ku2.rows[0].cells[i].text = h
    bg = '166534' if i == 0 else '991B1B'
    set_cell_bg(tbl_ku2.rows[0].cells[i], bg)
    for run in tbl_ku2.rows[0].cells[i].paragraphs[0].runs:
        run.font.bold  = True
        run.font.color.rgb = RGBColor(255,255,255)
        run.font.name  = 'Calibri'
        run.font.size  = Pt(10.5)

for r_i, (pro, con) in enumerate(ku2_data):
    row = tbl_ku2.rows[r_i + 1]
    row.cells[0].text = pro
    row.cells[1].text = con
    bg = 'F0FDF4' if r_i % 2 == 0 else 'FFFFFF'
    set_cell_bg(row.cells[0], bg)
    bg2 = 'FEF2F2' if r_i % 2 == 0 else 'FFFFFF'
    set_cell_bg(row.cells[1], bg2)
    for c_i in [0, 1]:
        for run in row.cells[c_i].paragraphs[0].runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(10)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════
# 5. PERBANDINGAN PAKET
# ══════════════════════════════════════════════════════════
add_section_title(doc, '5', 'PERBANDINGAN PAKET')

comp_headers = ['Parameter', 'Paket Starter (Tier 1)', 'Paket Profesional (Tier 2)']
comp_data = [
    ['vCPU',                   '4 Core',                    '6 Core'],
    ['RAM',                    '8 GB',                      '16 GB'],
    ['Storage',                '80 GB SSD NVMe',            '160 GB SSD NVMe'],
    ['Bandwidth',              '2 TB/bulan',                '4 TB/bulan'],
    ['Kapasitas Optimal',      's/d 300 siswa concurrent',  's/d 500+ siswa concurrent'],
    ['Uptime Garansi',         '99%',                       '99.5%'],
    ['Setup Fee',              'Rp 5.000.000',              'Rp 7.500.000'],
    ['Biaya Bulanan',          'Rp 1.500.000',              'Rp 2.500.000'],
    ['Total Tahun Pertama',    'Rp 23.000.000',             'Rp 37.500.000'],
    ['Tahun Selanjutnya',      'Rp 18.000.000/thn',         'Rp 30.000.000/thn'],
    ['Backup Mingguan',        '✅',                        '✅'],
    ['Monitoring 24/7',        '✅',                        '✅'],
    ['Training Admin',         '✅',                        '✅'],
    ['Cocok untuk',            'Tryout terjadwal bertahap', 'Tryout massal serentak'],
    ['Rekomendasi',            '—',                         '⭐ DIREKOMENDASIKAN'],
]

tbl_comp = doc.add_table(rows=1 + len(comp_data), cols=3)
tbl_comp.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_style(tbl_comp)

hdr_comp = tbl_comp.rows[0].cells
for i, h in enumerate(comp_headers):
    hdr_comp[i].text = h
    set_cell_bg(hdr_comp[i], '0F172A')
    for run in hdr_comp[i].paragraphs[0].runs:
        run.font.bold  = True
        run.font.color.rgb = RGBColor(255,255,255)
        run.font.name  = 'Calibri'
        run.font.size  = Pt(10.5)
        hdr_comp[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

for r_i, row_data in enumerate(comp_data):
    row = tbl_comp.rows[r_i + 1]
    for c_i, val in enumerate(row_data):
        row.cells[c_i].text = val
        if c_i == 0:
            set_cell_bg(row.cells[c_i], 'F1F5F9')
            for run in row.cells[c_i].paragraphs[0].runs:
                run.font.bold = True
        elif c_i == 1:
            bg = 'EFF6FF' if r_i % 2 == 0 else 'FFFFFF'
            set_cell_bg(row.cells[c_i], bg)
        else:
            bg = 'ECFDF5' if r_i % 2 == 0 else 'F0FDF4'
            set_cell_bg(row.cells[c_i], bg)
            if val == '⭐ DIREKOMENDASIKAN':
                for run in row.cells[c_i].paragraphs[0].runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(5, 150, 105)
        for run in row.cells[c_i].paragraphs[0].runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(10.5)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════
# 6. SKENARIO PENGGUNAAN
# ══════════════════════════════════════════════════════════
add_section_title(doc, '6', 'SKENARIO PENGGUNAAN')

skenario_items = [
    ('Skenario A — Tryout Bergilir (Cocok Tier 1)',
     'Tryout dibagi dalam 2–3 sesi berbeda waktu (misal pagi/siang/sore). '
     'Masing-masing sesi max 200 siswa. Server Tier 1 dapat menangani beban ini dengan nyaman.'),
    ('Skenario B — Tryout Massal Serentak (Wajib Tier 2)',
     '500 siswa login dan mengerjakan tryout pada waktu yang sama. '
     'Seluruh siswa submit dalam rentang 30 menit terakhir secara bersamaan. '
     'Tier 2 dirancang untuk skenario ini — performa tetap stabil.'),
    ('Skenario C — Pertumbuhan Jangka Panjang',
     'Jika jumlah siswa bertambah menjadi 700–800 dalam 1–2 tahun ke depan, '
     'Tier 2 masih mampu melayani. Upgrade ke server yang lebih besar dapat dilakukan '
     'tanpa mengubah arsitektur sistem.'),
]

for title_s, desc_s in skenario_items:
    p_t = doc.add_paragraph()
    r_t = p_t.add_run(f'▶  {title_s}')
    r_t.font.name  = 'Calibri'
    r_t.font.size  = Pt(11)
    r_t.font.bold  = True
    r_t.font.color.rgb = RGBColor(30, 64, 175)
    p_t.paragraph_format.space_before = Pt(8)
    p_t.paragraph_format.space_after  = Pt(2)

    p_d = doc.add_paragraph(desc_s)
    p_d.paragraph_format.left_indent = Cm(0.8)
    p_d.paragraph_format.space_after = Pt(4)
    for run in p_d.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(71, 85, 105)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════
# 7. SYARAT & KETENTUAN
# ══════════════════════════════════════════════════════════
add_section_title(doc, '7', 'SYARAT & KETENTUAN')

syarat_items = [
    'Kontrak minimum: 6 bulan untuk memastikan keberlangsungan layanan.',
    'Pembayaran setup fee dilakukan di muka sebelum deployment dimulai.',
    'Biaya bulanan dibayarkan setiap awal bulan (maksimal tanggal 5).',
    'Keterlambatan pembayaran lebih dari 15 hari dapat menyebabkan layanan ditangguhkan sementara.',
    'Data milik TRITON BIMBEL sepenuhnya — Atala Project tidak menggunakan data siswa untuk kepentingan lain.',
    'Jika terjadi gangguan di luar kontrol kami (force majeure, gangguan provider), tidak termasuk dalam SLA.',
    'Perubahan spesifikasi server (upgrade/downgrade) dapat dilakukan dengan pemberitahuan minimal 7 hari.',
    'Layanan mencakup satu (1) domain atau subdomain yang disepakati.',
]

for i, item in enumerate(syarat_items):
    p = doc.add_paragraph(f'{i+1}.  {item}')
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_after = Pt(3)
    for run in p.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(10.5)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════
# 8. REKOMENDASI AKHIR
# ══════════════════════════════════════════════════════════
add_section_title(doc, '8', 'REKOMENDASI AKHIR')

rek = doc.add_paragraph(
    'Kami merekomendasikan Paket Profesional (Tier 2) untuk TRITON BIMBEL. '
    'Dengan kapasitas 500 siswa concurrent, infrastruktur ini memberikan ruang yang cukup '
    'baik untuk operasional saat ini maupun pertumbuhan ke depan. Investasi yang sedikit lebih '
    'tinggi di awal akan menghindarkan risiko gangguan sistem saat tryout berlangsung — '
    'yang tentu berdampak langsung pada reputasi dan kepercayaan siswa.\n\n'
    'Kami siap untuk berdiskusi lebih lanjut, melakukan demo platform, atau menyesuaikan '
    'proposal sesuai kebutuhan spesifik TRITON BIMBEL.'
)
for run in rek.runs:
    run.font.name = 'Calibri'
    run.font.size = Pt(11)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════
# PENUTUP / TTD
# ══════════════════════════════════════════════════════════
div3 = doc.add_paragraph()
div_run3 = div3.add_run('─' * 85)
div_run3.font.color.rgb = RGBColor(30, 64, 175)
div_run3.font.size = Pt(9)

doc.add_paragraph()

ttd_tbl = doc.add_table(rows=1, cols=2)
ttd_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

left  = ttd_tbl.rows[0].cells[0]
right = ttd_tbl.rows[0].cells[1]

left_p  = left.paragraphs[0]
right_p = right.paragraphs[0]

def add_ttd(cell, title, company, name):
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.font.name = 'Calibri'
    r.font.size = Pt(10.5)
    r.font.bold = True

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(company)
    r2.font.name = 'Calibri'
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(100, 116, 139)

    for _ in range(4):
        p_space = cell.add_paragraph()
        p_space.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run('( ' + name + ' )')
    r3.font.name = 'Calibri'
    r3.font.size = Pt(10.5)
    r3.font.bold = True

    p4 = cell.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run('___________________')
    r4.font.name = 'Calibri'
    r4.font.size = Pt(10)

add_ttd(left,  'Hormat kami,',      'Atala Project',   'Pimpinan / PIC')
add_ttd(right, 'Menyetujui,',       'TRITON BIMBEL',   'Pimpinan / PIC')

doc.add_paragraph()
doc.add_paragraph()

footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_f = footer_p.add_run('Atala Project  |  Technology Solutions & Digital Infrastructure')
r_f.font.name  = 'Calibri'
r_f.font.size  = Pt(9)
r_f.font.color.rgb = RGBColor(100, 116, 139)

# ── Save
output = '/Users/indriregita/Desktop/ATALA PROJECT/Project/tritonapp/proposal/Proposal_Layanan_TRITON_BIMBEL.docx'
doc.save(output)
print(f'✅ Saved: {output}')
